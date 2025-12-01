"""
API para geração de Relatórios LSP-R
VERSÃO 2.3.0 - HTML Email com primeira página completa
"""

from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import subprocess
import shutil
import base64
from PyPDF2 import PdfMerger
import logging
import re

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = FastAPI(title="API Relatório LSP-R", version="2.3.0")

# Diretórios
BASE_DIR = Path(__file__).parent
TEMPLATES_DIR = BASE_DIR / "templates"
CORPOS_PDF_DIR = BASE_DIR / "assets" / "corpos_pdf"
TEMP_DIR = BASE_DIR / "temp"
TEMP_DIR.mkdir(exist_ok=True)

ARQUIVOS_VALIDOS = [
    'relatório_mais_ação_menos_mensagem',
    'relatório_mais_ação_menos_pessoas',
    'relatório_mais_ação_menos_tempo',
    'relatório_mais_mensagem_menos_ação',
    'relatório_mais_mensagem_menos_pessoas',
    'relatório_mais_mensagem_menos_tempo',
    'relatório_mais_pessoas_e_menos_ação',
    'relatório_mais_pessoas_e_menos_mensagem',
    'relatório_mais_pessoas_e_menos_tempo',
    'relatório_mais_tempo_e_menos_ação',
    'relatório_mais_tempo_e_menos_mensagem',
    'relatório_mais_tempo_e_menos_pessoas'
]

NOMES_ESTILOS = {
    "PESSOAS": "Pessoas (Relacional)",
    "ACAO": "Ação (Processo)",
    "TEMPO": "Tempo (Solução imediata)",
    "MENSAGEM": "Mensagem (Conteúdo / Analítico)"
}

NOMES_ESTILOS_LONGOS = {
    "PESSOAS": "Orientado para Pessoas (Relacional)",
    "ACAO": "Orientado para Ação (Processo)",
    "TEMPO": "Orientado para o Tempo (Solução imediata)",
    "MENSAGEM": "Orientado para Mensagem (Conteúdo / Analítico)"
}

class Pontuacoes(BaseModel):
    PESSOAS: int = Field(..., ge=0, le=60)
    ACAO: int = Field(..., ge=0, le=60)
    TEMPO: int = Field(..., ge=0, le=60)
    MENSAGEM: int = Field(..., ge=0, le=60)

class RelatorioRequest(BaseModel):
    participante: str = Field(..., min_length=1)
    pontuacoes: Pontuacoes
    predominante: str = Field(..., pattern="^(PESSOAS|ACAO|TEMPO|MENSAGEM)$")
    menosDesenvolvido: str = Field(..., pattern="^(PESSOAS|ACAO|TEMPO|MENSAGEM)$")
    arquivo: str


def remover_bordas_tabela(tabela):
    """Remove todas as bordas de uma tabela"""
    for row in tabela.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def criar_tabela_pontuacoes(doc, dados: RelatorioRequest):
    """Cria tabela DOCX real com bordas invisíveis"""
    logger.info("→ Criando tabela de pontuações...")
    
    # Criar tabela: 5 linhas (cabeçalho + 4 dados), 2 colunas
    tabela = doc.add_table(rows=5, cols=2)
    
    # Remover bordas
    remover_bordas_tabela(tabela)
    
    # CABEÇALHO
    cabecalho_cells = tabela.rows[0].cells
    cabecalho_cells[0].text = "Estilo de escuta"
    cabecalho_cells[1].text = "Pontuação"
    
    # Centralizar cabeçalho "Pontuação"
    cabecalho_cells[1].paragraphs[0].alignment = 1  # CENTER
    
    # DADOS (ordem correta)
    dados_tabela = [
        (NOMES_ESTILOS["PESSOAS"], str(dados.pontuacoes.PESSOAS)),
        (NOMES_ESTILOS["ACAO"], str(dados.pontuacoes.ACAO)),
        (NOMES_ESTILOS["TEMPO"], str(dados.pontuacoes.TEMPO)),
        (NOMES_ESTILOS["MENSAGEM"], str(dados.pontuacoes.MENSAGEM))
    ]
    
    for i, (nome, pontuacao) in enumerate(dados_tabela, start=1):
        row = tabela.rows[i]
        row.cells[0].text = nome
        row.cells[1].text = pontuacao
        
        # Alinhar número ao CENTRO
        row.cells[1].paragraphs[0].alignment = 1  # CENTER
        
        logger.info(f"  ✓ Linha {i}: {nome} = {pontuacao}")
    
    # Aplicar fonte DejaVu Sans em toda tabela
    for row in tabela.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'DejaVu Sans'
                    run.font.size = Pt(12)
    
    logger.info("✓ Tabela criada com sucesso")
    return tabela


def substituir_campos_docx(doc_path: Path, dados: RelatorioRequest, output_path: Path):
    """Substitui campos e cria tabela"""
    try:
        logger.info("="*60)
        logger.info(f"PROCESSANDO: {doc_path.name}")
        logger.info(f"Participante: {dados.participante}")
        logger.info("="*60)
        
        doc = Document(doc_path)
        logger.info(f"✓ Documento carregado: {len(doc.paragraphs)} parágrafos")
        
        paragrafos_para_remover = []
        paragrafo_tabela_idx = None
        
        # 1. SUBSTITUIR NOME e MARCAR LINHAS DA TABELA para remoção
        for i, para in enumerate(doc.paragraphs):
            texto = para.text
            
            # Substituir nome
            if "Nome completo" in texto:
                for run in para.runs:
                    if "Nome completo" in run.text:
                        run.text = run.text.replace("Nome completo", dados.participante)
                        logger.info(f"✓ Nome substituído: {dados.participante}")
            
            # Identificar cabeçalho da tabela
            if "Estilo de escuta" in texto and "Pontuação" in texto:
                logger.info(f"✓ Encontrado cabeçalho da tabela (parágrafo {i})")
                paragrafo_tabela_idx = i
                paragrafos_para_remover.append(i)
                continue
            
            # Marcar 4 linhas seguintes para remoção
            if paragrafo_tabela_idx is not None:
                offset = i - paragrafo_tabela_idx
                if 1 <= offset <= 4:
                    for estilo_nome in NOMES_ESTILOS.values():
                        if estilo_nome in texto:
                            logger.info(f"✓ Marcado para remoção: linha {i}")
                            paragrafos_para_remover.append(i)
                            break
            
            # Substituir predominante/menos desenvolvido
            if "Estilo predominante:" in texto:
                novo_texto = re.sub(
                    r'(Estilo predominante:\s*)(.+)',
                    r'\1' + NOMES_ESTILOS_LONGOS[dados.predominante],
                    "".join(run.text for run in para.runs)
                )
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = novo_texto
                logger.info("✓ Predominante substituído")
            
            if "Estilo menos desenvolvido:" in texto:
                novo_texto = re.sub(
                    r'(Estilo menos desenvolvido:\s*)(.+)',
                    r'\1' + NOMES_ESTILOS_LONGOS[dados.menosDesenvolvido],
                    "".join(run.text for run in para.runs)
                )
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = novo_texto
                logger.info("✓ Menos desenvolvido substituído")
        
        # 2. REMOVER PARÁGRAFOS DA TABELA ANTIGA
        logger.info(f"→ Removendo {len(paragrafos_para_remover)} parágrafos...")
        for idx in sorted(paragrafos_para_remover, reverse=True):
            p = doc.paragraphs[idx]._element
            p.getparent().remove(p)
        
        # 3. INSERIR TABELA DOCX REAL
        if paragrafo_tabela_idx is not None:
            logger.info("→ Inserindo tabela DOCX...")
            
            # Criar tabela
            tabela = criar_tabela_pontuacoes(doc, dados)
            
            # Inserir no documento (antes do próximo parágrafo)
            para_ref = doc.paragraphs[max(0, paragrafo_tabela_idx - 1)]._element
            para_ref.addnext(tabela._element)
            
            logger.info("✓ Tabela inserida")
        
        # 4. APLICAR DEJAVU SANS em todo documento
        logger.info("→ Aplicando DejaVu Sans 12pt...")
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.name = 'DejaVu Sans'
                run.font.size = Pt(12)
                run.font.highlight_color = None
        
        # Salvar
        doc.save(output_path)
        logger.info(f"✓ Documento salvo: {output_path.name}")
        logger.info("="*60)
        
        return True
        
    except Exception as e:
        logger.error(f"✗ ERRO: {e}", exc_info=True)
        raise


def converter_docx_para_pdf(docx_path: Path, pdf_path: Path):
    """Converte DOCX para PDF usando LibreOffice"""
    try:
        logger.info(f"→ Convertendo para PDF...")
        
        comando = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(pdf_path.parent),
            str(docx_path)
        ]
        
        result = subprocess.run(comando, capture_output=True, text=True, timeout=30)
        
        if result.returncode != 0:
            raise Exception(f"Conversão falhou: {result.stderr}")
        
        # Renomear se necessário
        arquivo_gerado = pdf_path.parent / f"{docx_path.stem}.pdf"
        if arquivo_gerado != pdf_path and arquivo_gerado.exists():
            shutil.move(str(arquivo_gerado), str(pdf_path))
        
        logger.info(f"✓ PDF gerado: {pdf_path.name}")
        return True
        
    except Exception as e:
        logger.error(f"✗ Erro na conversão: {e}")
        raise


def juntar_pdfs(capa_pdf: Path, corpo_pdf: Path, output_pdf: Path):
    """Junta capa e corpo em um PDF final"""
    try:
        logger.info("→ Juntando PDFs...")
        
        merger = PdfMerger()
        merger.append(str(capa_pdf))
        merger.append(str(corpo_pdf))
        merger.write(str(output_pdf))
        merger.close()
        
        logger.info(f"✓ PDF completo: {output_pdf.name}")
        return True
        
    except Exception as e:
        logger.error(f"✗ Erro ao juntar: {e}")
        raise


def gerar_html_capa(dados: RelatorioRequest) -> str:
    """Gera HTML com formatação IDÊNTICA ao documento Word"""
    dados_tabela = [
        ("Pessoas (Relacional)", dados.pontuacoes.PESSOAS),
        ("Ação (Processo)", dados.pontuacoes.ACAO),
        ("Tempo (Solução imediata)", dados.pontuacoes.TEMPO),
        ("Mensagem (Conteúdo / Analítico)", dados.pontuacoes.MENSAGEM)
    ]
    
    predominante = NOMES_ESTILOS_LONGOS[dados.predominante]
    menos_desenvolvido = NOMES_ESTILOS_LONGOS[dados.menosDesenvolvido]
    
    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: Aptos, Calibri, 'Segoe UI', Arial, sans-serif;
            font-size: 12pt;
            line-height: 1.15;
            color: #000000;
            background-color: #ffffff;
            padding: 40px 60px;
            max-width: 800px;
            margin: 0 auto;
        }}
        
        .logo-container {{
            text-align: center;
            margin-bottom: 20px;
        }}
        
        .logo {{
            width: 80px;
            height: auto;
        }}
        
        h1 {{
            font-size: 14pt;
            font-weight: bold;
            text-align: center;
            margin: 20px 0 30px 0;
            color: #000000;
        }}
        
        .participante {{
            font-size: 12pt;
            margin-bottom: 25px;
            line-height: 1.15;
        }}
        
        .participante strong {{
            font-weight: bold;
        }}
        
        h2 {{
            font-size: 12pt;
            font-weight: bold;
            margin: 25px 0 15px 0;
            color: #000000;
        }}
        
        .tabela-container {{
            margin: 15px 0 25px 0;
        }}
        
        .tabela-linha {{
            display: flex;
            margin-bottom: 2px;
        }}
        
        .tabela-cabecalho {{
            font-weight: bold;
            margin-bottom: 5px;
        }}
        
        .tabela-col-esquerda {{
            flex: 1;
            padding-right: 20px;
        }}
        
        .tabela-col-direita {{
            width: 100px;
            text-align: center;
        }}
        
        .destaque-box {{
            margin: 20px 0;
        }}
        
        .destaque-box p {{
            margin: 8px 0;
            line-height: 1.15;
        }}
        
        .destaque-box strong {{
            font-weight: bold;
        }}
        
        h3 {{
            font-size: 12pt;
            font-weight: bold;
            margin: 25px 0 12px 0;
            color: #000000;
        }}
        
        .descricao-paragrafo {{
            margin: 12px 0;
            text-indent: 0;
            line-height: 1.15;
        }}
        
        .descricao-paragrafo strong {{
            font-weight: bold;
        }}
        
        .footer-text {{
            margin-top: 25px;
            text-align: left;
            font-style: italic;
            line-height: 1.15;
        }}
        
        @media print {{
            body {{
                padding: 20px;
            }}
        }}
    </style>
</head>
<body>
    <div class="logo-container">
        <svg class="logo" width="80" height="80" viewBox="0 0 100 100" xmlns="http://www.w3.org/2000/svg">
            <circle cx="50" cy="50" r="45" fill="#4A90E2" opacity="0.2"/>
            <circle cx="30" cy="35" r="8" fill="#4A90E2"/>
            <circle cx="70" cy="35" r="8" fill="#4A90E2"/>
            <circle cx="50" cy="50" r="8" fill="#4A90E2"/>
            <circle cx="35" cy="60" r="6" fill="#4A90E2"/>
            <circle cx="65" cy="60" r="6" fill="#4A90E2"/>
            <circle cx="45" cy="70" r="5" fill="#4A90E2"/>
            <circle cx="55" cy="70" r="5" fill="#4A90E2"/>
            <path d="M30 35 L50 50 M70 35 L50 50 M50 50 L35 60 M50 50 L65 60 M35 60 L45 70 M65 60 L55 70" 
                  stroke="#4A90E2" stroke-width="2" fill="none"/>
        </svg>
    </div>

    <h1>Relatório de Perfil de Escuta e Comunicação</h1>
    
    <p class="participante"><strong>Participante:</strong> {dados.participante}</p>
    
    <h2>Resultado geral</h2>
    
    <div class="tabela-container">
        <div class="tabela-linha tabela-cabecalho">
            <div class="tabela-col-esquerda">Estilo de escuta</div>
            <div class="tabela-col-direita">Pontuação</div>
        </div>"""
    
    for nome, pont in dados_tabela:
        html += f"""
        <div class="tabela-linha">
            <div class="tabela-col-esquerda">{nome}</div>
            <div class="tabela-col-direita">{pont}</div>
        </div>"""
    
    html += f"""
    </div>
    
    <div class="destaque-box">
        <p><strong>Estilo predominante:</strong> {predominante}</p>
        <p><strong>Estilo menos desenvolvido:</strong> {menos_desenvolvido}</p>
    </div>
    
    <h3>Descrição geral dos 4 estilos:</h3>
    
    <p class="descricao-paragrafo">
        <strong>Orientado para Pessoas (Relacional):</strong> valoriza o vínculo e empatia. Escuta com atenção às emoções e constrói confiança pela proximidade.
    </p>
    
    <p class="descricao-paragrafo">
        <strong>Orientado para Ação (Processo):</strong> prefere conversas diretas, voltadas à solução e ao resultado. Gosta de foco e clareza, mas pode soar apressado.
    </p>
    
    <p class="descricao-paragrafo">
        <strong>Orientado para o Tempo (Solução imediata):</strong> preza pela objetividade e gosta de ritmo na conversa. Evita desvios e busca eficiência.
    </p>
    
    <p class="descricao-paragrafo">
        <strong>Orientado para Mensagem (Conteúdo / Analítico):</strong> escuta para compreender o sentido exato do que está sendo dito. Avalia argumentos, identifica contradições e busca precisão na comunicação.
    </p>
    
    <p class="footer-text">
        Essas informações serão aprofundadas no relatório anexo.
    </p>
</body>
</html>"""
    
    return html


@app.get("/")
async def root():
    return {"message": "API Relatório LSP-R", "version": "2.3.0"}


@app.get("/health")
async def health():
    checks = {
        "templates_dir": TEMPLATES_DIR.exists(),
        "corpos_pdf_dir": CORPOS_PDF_DIR.exists(),
        "libreoffice": shutil.which("libreoffice") is not None
    }
    return {
        "status": "ok" if all(checks.values()) else "warning",
        "version": "2.3.0",
        "checks": checks
    }


@app.get("/templates-disponiveis")
async def listar_templates():
    templates_completos = []
    for arquivo in ARQUIVOS_VALIDOS:
        docx_exists = (TEMPLATES_DIR / f"{arquivo}.docx").exists()
        pdf_exists = (CORPOS_PDF_DIR / f"{arquivo}.pdf").exists()
        if docx_exists and pdf_exists:
            templates_completos.append(arquivo)
    
    return {"templates_completos": templates_completos, "total": len(templates_completos)}


@app.post("/gerar-html-email")
async def gerar_html_email(dados: RelatorioRequest):
    """Retorna apenas HTML para corpo de email"""
    try:
        if dados.predominante == dados.menosDesenvolvido:
            raise HTTPException(400, "Predominante e menos desenvolvido não podem ser iguais")
        
        html = gerar_html_capa(dados)
        
        return {
            "html": html,
            "participante": dados.participante,
            "version": "2.3.0"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erro ao gerar HTML: {e}")
        raise HTTPException(500, str(e))


@app.post("/gerar-relatorio-completo")
async def gerar_relatorio_completo(dados: RelatorioRequest):
    """Gera PDF E HTML em uma única chamada"""
    try:
        logger.info("REQUISIÇÃO COMPLETA (PDF + HTML)")
        
        if dados.predominante == dados.menosDesenvolvido:
            raise HTTPException(400, "Estilos não podem ser iguais")
        
        if dados.arquivo not in ARQUIVOS_VALIDOS:
            raise HTTPException(400, "Arquivo inválido")
        
        template_docx = TEMPLATES_DIR / f"{dados.arquivo}.docx"
        corpo_pdf = CORPOS_PDF_DIR / f"{dados.arquivo}.pdf"
        
        if not template_docx.exists() or not corpo_pdf.exists():
            raise HTTPException(404, "Template ou corpo não encontrado")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_docx = TEMP_DIR / f"capa_{timestamp}.docx"
        temp_pdf = TEMP_DIR / f"capa_{timestamp}.pdf"
        temp_final = TEMP_DIR / f"final_{timestamp}.pdf"
        
        substituir_campos_docx(template_docx, dados, temp_docx)
        converter_docx_para_pdf(temp_docx, temp_pdf)
        juntar_pdfs(temp_pdf, corpo_pdf, temp_final)
        
        with open(temp_final, 'rb') as f:
            pdf_base64 = base64.b64encode(f.read()).decode('utf-8')
        
        html = gerar_html_capa(dados)
        
        logger.info("✓ PDF e HTML gerados")
        
        return {
            "success": True,
            "pdf_base64": pdf_base64,
            "html": html,
            "filename": f"relatorio_{dados.participante.replace(' ', '_')}.pdf",
            "participante": dados.participante
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erro: {e}", exc_info=True)
        raise HTTPException(500, str(e))


@app.post("/gerar-relatorio")
async def gerar_relatorio(dados: RelatorioRequest):
    """Gera PDF e retorna arquivo para download"""
    try:
        logger.info("="*60)
        logger.info("NOVA REQUISIÇÃO")
        logger.info("="*60)
        
        if dados.predominante == dados.menosDesenvolvido:
            raise HTTPException(400, "Predominante e menos desenvolvido não podem ser iguais")
        
        if dados.arquivo not in ARQUIVOS_VALIDOS:
            raise HTTPException(400, f"Arquivo inválido: {dados.arquivo}")
        
        template_docx = TEMPLATES_DIR / f"{dados.arquivo}.docx"
        corpo_pdf = CORPOS_PDF_DIR / f"{dados.arquivo}.pdf"
        
        if not template_docx.exists():
            raise HTTPException(404, f"Template não encontrado")
        
        if not corpo_pdf.exists():
            raise HTTPException(404, f"Corpo não encontrado")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_docx = TEMP_DIR / f"capa_{timestamp}.docx"
        temp_pdf = TEMP_DIR / f"capa_{timestamp}.pdf"
        temp_final = TEMP_DIR / f"final_{timestamp}.pdf"
        
        try:
            substituir_campos_docx(template_docx, dados, temp_docx)
            converter_docx_para_pdf(temp_docx, temp_pdf)
            juntar_pdfs(temp_pdf, corpo_pdf, temp_final)
            
            logger.info("="*60)
            logger.info("✓✓✓ SUCESSO ✓✓✓")
            logger.info("="*60)
            
            return FileResponse(
                path=str(temp_final),
                media_type="application/pdf",
                filename=f"relatorio_{dados.participante.replace(' ', '_')}.pdf"
            )
            
        finally:
            pass
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"✗✗✗ ERRO FATAL: {e}", exc_info=True)
        raise HTTPException(500, str(e))


@app.on_event("startup")
async def startup():
    logger.info("="*60)
    logger.info("API Relatório LSP-R v2.3.0")
    logger.info("="*60)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=3344)
